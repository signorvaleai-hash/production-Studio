from __future__ import annotations

from html import escape
from pathlib import Path
from uuid import uuid4

from docx import Document
from docx.shared import Inches, Pt
from ebooklib import epub
from reportlab.lib.pagesizes import inch
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch as rl_inch
from reportlab.lib.utils import ImageReader
from reportlab.platypus import Image, PageBreak, Paragraph, SimpleDocTemplate, Spacer

from formatter.parser import Block, Manuscript
from formatter.profiles import FormatProfile


def export_docx(
    manuscript: Manuscript,
    output_path: Path,
    profile: FormatProfile,
    front_cover: Path | None = None,
    back_cover: Path | None = None,
) -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(profile.page_width_in)
    section.page_height = Inches(profile.page_height_in)
    section.left_margin = Inches(profile.margin_in)
    section.right_margin = Inches(profile.margin_in)
    section.top_margin = Inches(profile.margin_in)
    section.bottom_margin = Inches(profile.margin_in)

    normal = doc.styles["Normal"]
    normal.font.name = profile.body_font
    normal.font.size = Pt(profile.body_size_pt)
    normal.paragraph_format.first_line_indent = Inches(profile.first_line_indent_in)
    normal.paragraph_format.line_spacing = profile.line_spacing

    heading = doc.styles["Heading 1"]
    heading.font.name = profile.body_font
    heading.font.size = Pt(profile.body_size_pt + 4)

    if front_cover:
        doc.add_picture(str(front_cover), width=Inches(profile.page_width_in - 2 * profile.margin_in))
        doc.add_page_break()

    title_para = doc.add_paragraph(manuscript.title)
    title_para.style = "Title"
    doc.add_paragraph(f"By {manuscript.author}")
    doc.add_page_break()

    wrote_any_chapter = False
    for block in manuscript.blocks:
        if block.kind == "chapter":
            if wrote_any_chapter:
                doc.add_page_break()
            h = doc.add_paragraph(block.text)
            h.style = "Heading 1"
            wrote_any_chapter = True
        else:
            doc.add_paragraph(block.text)

    if back_cover:
        doc.add_page_break()
        doc.add_picture(str(back_cover), width=Inches(profile.page_width_in - 2 * profile.margin_in))

    doc.save(output_path)


def export_epub(
    manuscript: Manuscript,
    output_path: Path,
    profile: FormatProfile,
    front_cover: Path | None = None,
    back_cover: Path | None = None,
) -> None:
    book = epub.EpubBook()
    book.set_identifier(str(uuid4()))
    book.set_title(manuscript.title)
    book.set_language("en")
    if manuscript.author:
        book.add_author(manuscript.author)

    style_content = f"""
    body {{
        font-family: '{profile.body_font}', serif;
        line-height: {profile.line_spacing};
        margin: 0;
        padding: 0;
    }}
    h1 {{
        text-align: center;
        margin-top: 2.2em;
        margin-bottom: 1.2em;
    }}
    p {{
        text-indent: {profile.first_line_indent_in}in;
        margin: 0 0 0.6em 0;
    }}
    .title-page {{
        text-align: center;
        margin-top: 35%;
    }}
    .cover {{
        text-align: center;
    }}
    img {{
        max-width: 100%;
        height: auto;
    }}
    """
    nav_css = epub.EpubItem(
        uid="style_nav",
        file_name="style/nav.css",
        media_type="text/css",
        content=style_content.encode("utf-8"),
    )
    book.add_item(nav_css)

    chapters: list[epub.EpubHtml] = []

    title_page = epub.EpubHtml(title="Title Page", file_name="title.xhtml", lang="en")
    title_page.content = (
        f"<html><body><div class='title-page'><h1>{escape(manuscript.title)}</h1>"
        f"<p>By {escape(manuscript.author or 'Unknown Author')}</p></div></body></html>"
    )
    book.add_item(title_page)
    chapters.append(title_page)

    if front_cover:
        cover_bytes = front_cover.read_bytes()
        book.set_cover(front_cover.name, cover_bytes)
        front_image_item = epub.EpubImage(
            uid="front_cover_img",
            file_name=f"images/{front_cover.name}",
            media_type=_guess_image_media_type(front_cover),
            content=cover_bytes,
        )
        book.add_item(front_image_item)
        cover_page = epub.EpubHtml(title="Cover", file_name="front-cover.xhtml", lang="en")
        cover_page.content = (
            f"<html><body><div class='cover'><img src='images/{escape(front_cover.name)}' alt='Front cover'/></div>"
            "</body></html>"
        )
        book.add_item(cover_page)
        chapters.insert(0, cover_page)

    chapter_buckets = _group_blocks_by_chapter(manuscript.blocks)
    for idx, (chapter_title, paragraphs) in enumerate(chapter_buckets, start=1):
        chapter = epub.EpubHtml(title=chapter_title, file_name=f"chap_{idx}.xhtml", lang="en")
        body_parts = [f"<h1>{escape(chapter_title)}</h1>"] if chapter_title else []
        body_parts.extend(f"<p>{escape(p)}</p>" for p in paragraphs)
        chapter.content = f"<html><body>{''.join(body_parts)}</body></html>"
        book.add_item(chapter)
        chapters.append(chapter)

    if back_cover:
        back_bytes = back_cover.read_bytes()
        back_item = epub.EpubImage(
            uid="back_cover_img",
            file_name=f"images/{back_cover.name}",
            media_type=_guess_image_media_type(back_cover),
            content=back_bytes,
        )
        book.add_item(back_item)
        back_page = epub.EpubHtml(title="Back Cover", file_name="back-cover.xhtml", lang="en")
        back_page.content = (
            f"<html><body><div class='cover'><img src='images/{escape(back_cover.name)}' alt='Back cover'/></div>"
            "</body></html>"
        )
        book.add_item(back_page)
        chapters.append(back_page)

    book.toc = tuple(chapters)
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    book.spine = ["nav", *chapters]
    epub.write_epub(str(output_path), book)


def export_pdf(
    manuscript: Manuscript,
    output_path: Path,
    profile: FormatProfile,
    front_cover: Path | None = None,
    back_cover: Path | None = None,
) -> None:
    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=(profile.page_width_in * inch, profile.page_height_in * inch),
        leftMargin=profile.margin_in * inch,
        rightMargin=profile.margin_in * inch,
        topMargin=profile.margin_in * inch,
        bottomMargin=profile.margin_in * inch,
    )

    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        name="Body",
        parent=styles["BodyText"],
        fontName=profile.body_font,
        fontSize=profile.body_size_pt,
        leading=profile.body_size_pt * profile.line_spacing,
        firstLineIndent=profile.first_line_indent_in * rl_inch,
        spaceAfter=8,
    )
    chapter_style = ParagraphStyle(
        name="Chapter",
        parent=styles["Heading1"],
        fontName=profile.body_font,
        fontSize=profile.body_size_pt + 4,
        leading=(profile.body_size_pt + 4) * 1.2,
        alignment=1,  # center
        spaceBefore=20,
        spaceAfter=20,
    )

    story = []

    if front_cover:
        story.append(_build_pdf_cover_image(front_cover, profile))
        story.append(PageBreak())

    story.append(Paragraph(escape(manuscript.title), styles["Title"]))
    story.append(Spacer(1, 12))
    story.append(Paragraph(escape(f"By {manuscript.author or 'Unknown Author'}"), styles["Heading3"]))
    story.append(PageBreak())

    wrote_chapter = False
    for block in manuscript.blocks:
        if block.kind == "chapter":
            if wrote_chapter:
                story.append(PageBreak())
            story.append(Paragraph(escape(block.text), chapter_style))
            wrote_chapter = True
        else:
            story.append(Paragraph(escape(block.text), body))

    if back_cover:
        story.append(PageBreak())
        story.append(_build_pdf_cover_image(back_cover, profile))

    doc.build(story)


def _group_blocks_by_chapter(blocks: list[Block]) -> list[tuple[str, list[str]]]:
    buckets: list[tuple[str, list[str]]] = []
    current_title = "Chapter 1"
    current_paragraphs: list[str] = []
    seen_chapter = False

    for block in blocks:
        if block.kind == "chapter":
            if current_paragraphs or seen_chapter:
                buckets.append((current_title, current_paragraphs))
            current_title = block.text
            current_paragraphs = []
            seen_chapter = True
        else:
            current_paragraphs.append(block.text)

    if current_paragraphs or not buckets:
        buckets.append((current_title, current_paragraphs))

    return buckets


def _guess_image_media_type(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".jpg", ".jpeg"}:
        return "image/jpeg"
    if suffix == ".png":
        return "image/png"
    if suffix == ".webp":
        return "image/webp"
    if suffix in {".tif", ".tiff"}:
        return "image/tiff"
    return "application/octet-stream"


def _build_pdf_cover_image(path: Path, profile: FormatProfile) -> Image:
    frame_w = profile.page_width_in * inch - (2 * profile.margin_in * inch)
    frame_h = profile.page_height_in * inch - (2 * profile.margin_in * inch)
    # Small safety margin avoids frame rounding errors in reportlab layout.
    max_w = frame_w * 0.98
    max_h = frame_h * 0.98

    try:
        pixel_w, pixel_h = ImageReader(str(path)).getSize()
        scale = min(max_w / pixel_w, max_h / pixel_h)
        img = Image(str(path), width=pixel_w * scale, height=pixel_h * scale)
    except Exception:
        img = Image(str(path))
        img._restrictSize(max_w, max_h)

    img.hAlign = "CENTER"
    return img
